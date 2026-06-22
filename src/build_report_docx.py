from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from pygments import highlight
from pygments.formatters import ImageFormatter
from pygments.lexers import PythonLexer


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "KIE4031_Stock_Price_Prediction_Report.docx"
FRONT_PAGE_LOGO = ROOT / "assets" / "front_page_logo.png"
PREPROCESSING_VISUAL = ROOT / "assets" / "preprocessing_flowchart.png"
RNN_COMPARISON_VISUAL = ROOT / "RNN_GRU_LSTM_comp.png"
CODE_IMAGE_DIR = ROOT / "assets" / "code_snippets"
FONT_NAME = "Times New Roman"
BLACK = (0, 0, 0)
SOURCE_CODE_URL = "https://github.com/JoeVey22/ML_Summative_Assessment"


def set_run(run, size=11, bold=False, color=BLACK):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run(run, size=16 if level == 1 else 13, bold=True)
    document._next_body_paragraph_no_indent = True
    return paragraph


def add_blank_paragraph(document, centered=False):
    paragraph = document.add_paragraph()
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def add_centered_cover_line(document, text, size=16, bold=True):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold)
    return paragraph


def set_table_borders_none(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_front_page(document):
    for _ in range(3):
        add_blank_paragraph(document)

    if FRONT_PAGE_LOGO.exists():
        document.add_picture(str(FRONT_PAGE_LOGO), width=Inches(5.25))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        add_blank_paragraph(document)

    add_centered_cover_line(document, "KIE4031: MACHINE LEARNING")
    add_centered_cover_line(document, "SEMESTER 2, 2025/2026")
    add_blank_paragraph(document, centered=True)
    add_centered_cover_line(document, "FINAL SUMMATIVE ASSESSMENT")
    add_blank_paragraph(document, centered=True)
    add_centered_cover_line(document, "STOCK PRICE PREDICTION USING RNN-BASED MODELS")
    add_blank_paragraph(document, centered=True)

    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders_none(table)

    rows = [
        ("NAME:", "KHOO JOE VEY"),
        ("MATRIC NUMBER:", "22004672/1"),
        ("LECTURER:", "ASSOCIATE PROF. IR. DR. ANIS SALWA BINTI MOHD KHAIRUDDIN"),
    ]
    for row_idx, (label, value) in enumerate(rows):
        row = table.rows[row_idx]
        set_cell_width(row.cells[0], 2340)
        set_cell_width(row.cells[1], 5575)
        for cell, text in zip(row.cells, (label, value)):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            set_run(run, size=12)

    document.add_page_break()


def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    if getattr(document, "_next_body_paragraph_no_indent", False):
        paragraph.paragraph_format.first_line_indent = None
        document._next_body_paragraph_no_indent = False
    else:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_source_code_link(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)

    label = paragraph.add_run("GitHub repository: ")
    set_run(label)

    relationship_id = document.part.relate_to(
        SOURCE_CODE_URL,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    link_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((color, underline))
    link_run.append(run_properties)
    link_text = OxmlElement("w:t")
    link_text.text = SOURCE_CODE_URL
    link_run.append(link_text)
    hyperlink.append(link_run)
    paragraph._p.append(hyperlink)

    add_paragraph(document, "Main training source file: src/train_stock_rnn.py")


def add_code_block(document, code):
    block_number = getattr(document, "_code_block_number", 0) + 1
    document._code_block_number = block_number
    CODE_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    image_path = CODE_IMAGE_DIR / f"code_snippet_{block_number}.png"

    png_bytes = highlight(
        code,
        PythonLexer(),
        ImageFormatter(
            style="friendly",
            font_name="Consolas",
            font_size=20,
            line_numbers=False,
            image_pad=24,
            line_pad=3,
            bg="#F7F7F7",
        ),
    )
    image_path.write_bytes(png_bytes)
    with Image.open(image_path) as image:
        image.save(image_path, dpi=(192, 192))
        width_inches = min(6.1, image.width / 192)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run(run)
    return paragraph


def add_figure(document, path, caption):
    if path.exists():
        document.add_picture(str(path), width=Inches(6.4))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_paragraph.add_run(caption)
        set_run(run, size=9, bold=False)


def add_table_caption(document, caption):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(caption)
    set_run(run, size=10, bold=True)
    return paragraph


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        if style_name in styles:
            styles[style_name].font.name = FONT_NAME
            styles[style_name].font.size = Pt(11)
            styles[style_name].font.color.rgb = RGBColor(*BLACK)

    add_front_page(document)

    add_heading(document, "1. Data Collection and Preprocessing")
    add_paragraph(
        document,
        "This project predicts Alphabet Inc. Class A stock prices using historical daily price data for GOOGL. "
        "The dataset was obtained from the public Yahoo Finance chart data endpoint through the Python "
        "script in src/train_stock_rnn.py. The final downloaded dataset covers 2019-01-02 to 2025-12-30 "
        "and contains 1,759 cleaned daily observations.",
    )
    add_paragraph(
        document,
        "The target variable is adjusted closing price because it accounts for corporate actions such as "
        "stock splits and dividends more consistently than raw close price. The raw data includes Date, "
        "Open, High, Low, Close, Adjusted Close, and Volume.",
    )
    for item in [
        "Data cleaning: The data was sorted by date to preserve the time-series order. Duplicate dates were removed so that each trading day appeared only once. Rows with missing adjusted close prices were removed because adjusted close is the prediction target. Invalid rows with non-positive adjusted close prices were also removed.",
        "Train/validation/test split: The cleaned dataset was split chronologically into 70% training data, 10% validation data, and 20% testing data. A chronological split was used because stock prediction is a time-series problem and future data must not be leaked into the training stage.",
        "Normalization: The adjusted close price was normalized using MinMaxScaler. The scaler was fitted only on the training set, then applied to the validation and testing data to prevent later-period information from influencing training.",
        "Sequence conversion: The normalized price series was converted into supervised learning sequences using a 60-day lookback window. The previous 60 adjusted closing prices were used as input features, and the next trading day's adjusted close price was used as the target output.",
    ]:
        add_number(document, item)
    add_paragraph(
        document,
        "The overall preprocessing workflow is summarized in Figure 1. It shows how the raw stock table is "
        "transformed into leakage-safe RNN input sequences.",
    )
    add_figure(document, PREPROCESSING_VISUAL, "Figure 1. Before/after view and flowchart of the preprocessing pipeline.")
    add_paragraph(
        document,
        "In the source code, the load_and_clean_data() function performs sorting, duplicate removal, missing-value "
        "removal, and invalid-price filtering. The MinMaxScaler section performs normalization, the validation "
        "set monitors overfitting during training, and the build_sequences() function converts the cleaned price "
        "data into RNN-ready input sequences.",
    )
    add_paragraph(document, "The following preprocessing code shows how the dataset was cleaned before training:")
    add_code_block(
        document,
        'def load_and_clean_data():\n'
        '    if not DATA_PATH.exists():\n'
        '        download_yahoo_chart(SYMBOL, START_DATE, END_DATE)\n\n'
        '    df = pd.read_csv(DATA_PATH, parse_dates=["Date"])\n'
        '    df = df.sort_values("Date").drop_duplicates("Date")\n'
        '    df = df.dropna(subset=["Adj Close"])\n'
        '    df = df[df["Adj Close"] > 0].reset_index(drop=True)\n'
        '    return df',
    )
    add_paragraph(document, "The chronological split and scaler fitting were also kept separate to reduce data leakage:")
    add_code_block(
        document,
        'train_end_index = int(len(df) * (1 - VALIDATION_RATIO - TEST_RATIO))\n'
        'validation_end_index = int(len(df) * (1 - TEST_RATIO))\n'
        'train_close = df.loc[: train_end_index - 1, ["Adj Close"]]\n\n'
        'scaler = MinMaxScaler()\n'
        'scaler.fit(train_close)\n'
        'scaled_close = scaler.transform(df[["Adj Close"]])',
    )
    add_paragraph(document, "The cleaned and scaled series was converted into 60-day input windows using this sequence-building function:")
    add_code_block(
        document,
        'def build_sequences(values, lookback):\n'
        '    x, y = [], []\n'
        '    for idx in range(lookback, len(values)):\n'
        '        x.append(values[idx - lookback : idx])\n'
        '        y.append(values[idx])\n'
        '    return np.array(x, dtype=np.float32), np.array(y, dtype=np.float32)',
    )

    add_heading(document, "2. Investigation of the Proposed Technique")
    add_paragraph(
        document,
        "The proposed machine learning technique is an RNN-based time-series forecasting model. RNNs are "
        "designed for sequential data because they process observations in order and maintain hidden states "
        "that represent information from previous time steps. This makes them suitable for stock price "
        "prediction, where recent and medium-term historical prices can influence the next predicted value.",
    )
    add_paragraph(
        document,
        "Three RNN-based models were investigated: a vanilla RNN, a Gated Recurrent Unit model, and a Long "
        "Short-Term Memory model. A vanilla RNN is the simplest recurrent model and updates its hidden state "
        "at each time step. However, it can struggle with long-term dependencies because gradients may vanish "
        "or explode during backpropagation through time.",
    )
    add_paragraph(
        document,
        "GRU improves the standard RNN by using update and reset gates. These gates help the model decide how "
        "much past information should be kept and how much new information should be added. GRU is usually "
        "simpler than LSTM because it has fewer gates, so it can train faster while still handling sequential "
        "dependencies effectively.",
    )
    add_paragraph(
        document,
        "LSTM improves the standard RNN by using input, forget, and output gates with a memory cell. These gates "
        "control how information is stored, updated, and forgotten. LSTM is powerful for longer sequences, but "
        "it has more parameters than GRU and may require more data or tuning to perform best.",
    )
    for item in [
        "Input size: 1 feature, the normalized adjusted close price.",
        "Lookback window: 60 trading days.",
        "Two recurrent layers with hidden size 64.",
        "Dropout: 0.20 to reduce overfitting.",
        "Adam optimizer with learning rate 0.001 and mean squared error loss.",
        "Training duration: 40 epochs.",
        "Validation loss monitoring, with the best-validation epoch restored before final test evaluation.",
    ]:
        add_bullet(document, item)
    add_paragraph(
        document,
        "After comparison, GRU was selected as the best proposed model for this dataset because it achieved "
        "the lowest test error among the RNN-based models.",
    )
    add_paragraph(
        document,
        "Figure 2 compares the internal structures of RNN, LSTM, and GRU. It shows that the basic RNN mainly "
        "transforms the current input and previous hidden state through a tanh layer, while LSTM and GRU add "
        "gating operations using sigmoid functions, pointwise multiplication, and pointwise addition to control "
        "how much information is kept or forgotten.",
    )
    add_figure(document, RNN_COMPARISON_VISUAL, "Figure 2. Internal gate-level comparison of RNN, LSTM, and GRU.")
    add_paragraph(
        document,
        "The shared PyTorch model class made the comparison fair because the recurrent cell type was the main "
        "component changed between RNN, GRU, and LSTM:",
    )
    add_code_block(
        document,
        'class StockRecurrentModel(nn.Module):\n'
        '    def __init__(self, cell_type, input_size=1, hidden_size=64,\n'
        '                 num_layers=2, dropout=0.20):\n'
        '        super().__init__()\n'
        '        recurrent_layers = {"RNN": nn.RNN, "GRU": nn.GRU, "LSTM": nn.LSTM}\n'
        '        self.recurrent = recurrent_layers[cell_type](\n'
        '            input_size=input_size,\n'
        '            hidden_size=hidden_size,\n'
        '            num_layers=num_layers,\n'
        '            batch_first=True,\n'
        '            dropout=dropout,\n'
        '        )\n'
        '        self.regressor = nn.Sequential(\n'
        '            nn.Linear(hidden_size, 32), nn.ReLU(), nn.Linear(32, 1)\n'
        '        )\n\n'
        '    def forward(self, x):\n'
        '        output, _ = self.recurrent(x)\n'
        '        return self.regressor(output[:, -1, :])',
    )
    add_heading(document, "Comparison of RNN, GRU, and LSTM", level=2)
    add_table_caption(document, "Table 1. Comparison of RNN, GRU, and LSTM models.")
    comparison_table = document.add_table(rows=1, cols=3)
    comparison_table.style = "Table Grid"
    headers = ["Model", "Good points", "Bad points"]
    for idx, text in enumerate(headers):
        run = comparison_table.rows[0].cells[idx].paragraphs[0].add_run(text)
        set_run(run, bold=True)
    comparison_rows = [
        [
            "RNN",
            "Simple structure, fast to train, and easy to understand. It can model short-term sequential patterns and is useful as a basic recurrent baseline.",
            "Struggles with long-term dependencies because of vanishing or exploding gradients. It may forget older information in the 60-day sequence and usually performs worse than gated models.",
        ],
        [
            "GRU",
            "Uses update and reset gates, so it handles sequence memory better than a basic RNN. It has fewer parameters than LSTM, trains efficiently, and achieved the best result in this experiment.",
            "Less expressive than LSTM in some complex long-sequence problems. It can still overfit and still depends heavily on the quality of historical price data.",
        ],
        [
            "LSTM",
            "Uses input, forget, and output gates with a memory cell, making it strong for longer-term dependencies. It is widely used for time-series prediction and can model complex temporal patterns.",
            "More complex than RNN and GRU, with more parameters and higher training cost. On this single-feature GOOGL dataset, it performed worse than GRU, possibly because the extra complexity was not needed.",
        ],
    ]
    for row_values in comparison_rows:
        row = comparison_table.add_row()
        for idx, text in enumerate(row_values):
            run = row.cells[idx].paragraphs[0].add_run(text)
            set_run(run, size=9 if idx else 10, bold=idx == 0)

    add_heading(document, "3. Model Development, Evaluation, and Visualization")
    add_paragraph(
        document,
        "The model was implemented using Python and PyTorch. Pandas and NumPy were used for data handling, "
        "scikit-learn was used for scaling and evaluation metrics, and Matplotlib was used for visualization. "
        "The RNN, GRU, and LSTM models were compared with a 60-day moving average baseline.",
    )
    add_paragraph(document, "The training loop used mini-batch gradient descent, mean squared error loss, the Adam optimizer, and validation-loss monitoring:")
    add_code_block(
        document,
        'model = StockRecurrentModel(model_name).to(device)\n'
        'loss_fn = nn.MSELoss()\n'
        'optimizer = torch.optim.Adam(model.parameters(), lr=LEARNING_RATE)\n'
        'best_val_loss = float("inf")\n\n'
        'for epoch in range(EPOCHS):\n'
        '    model.train()\n'
        '    for batch_x, batch_y in train_loader:\n'
        '        batch_x = batch_x.to(device)\n'
        '        batch_y = batch_y.to(device)\n'
        '        optimizer.zero_grad()\n'
        '        prediction = model(batch_x)\n'
        '        loss = loss_fn(prediction, batch_y)\n'
        '        loss.backward()\n'
        '        optimizer.step()\n\n'
        '    model.eval()\n'
        '    with torch.no_grad():\n'
        '        val_loss = loss_fn(model(val_x), val_y)\n'
        '        if val_loss.item() < best_val_loss:\n'
        '            best_val_loss = val_loss.item()\n'
        '            best_state = deepcopy(model.state_dict())',
    )
    add_paragraph(document, "The evaluation function calculated the same four metrics used in the results table:")
    add_code_block(
        document,
        'def evaluate(y_true, y_pred):\n'
        '    rmse = math.sqrt(mean_squared_error(y_true, y_pred))\n'
        '    mae = mean_absolute_error(y_true, y_pred)\n'
        '    mape = np.mean(np.abs((y_true - y_pred) / y_true)) * 100\n'
        '    r2 = r2_score(y_true, y_pred)\n'
        '    return {"RMSE": rmse, "MAE": mae, "MAPE": mape, "R2": r2}',
    )

    add_table_caption(document, "Table 2. Model evaluation metrics on the test set.")
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Model", "RMSE", "MAE", "MAPE", "R2"]
    for idx, text in enumerate(headers):
        run = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        set_run(run, bold=True)
    rows = [
        ["RNN", "34.83", "21.53", "8.78%", "0.456"],
        ["GRU", "26.06", "15.54", "6.28%", "0.695"],
        ["LSTM", "34.17", "20.63", "8.37%", "0.476"],
        ["60-day Moving Average", "23.61", "19.22", "9.12%", "0.750"],
    ]
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            run = row.cells[idx].paragraphs[0].add_run(text)
            set_run(run)

    add_paragraph(
        document,
        "After adding validation monitoring, GRU remained the strongest RNN-based model. It achieved the lowest "
        "MAE and MAPE among all compared methods, but the 60-day moving average baseline achieved the lowest "
        "RMSE and highest R2. This indicates that GRU reduced average absolute percentage error well, while the "
        "simple baseline was still competitive on squared-error-based metrics.",
    )
    add_figure(document, ROOT / "outputs" / "stock_price_history.png", "Figure 3. GOOGL adjusted close price history.")
    add_figure(document, ROOT / "outputs" / "prediction_vs_actual.png", "Figure 4. Actual vs predicted test-set prices.")
    add_figure(document, ROOT / "outputs" / "training_loss.png", "Figure 5. RNN model training loss.")
    add_figure(document, ROOT / "outputs" / "validation_loss.png", "Figure 6. RNN model validation loss.")
    add_figure(document, ROOT / "outputs" / "model_rmse_comparison.png", "Figure 7. RMSE comparison across models.")

    add_heading(document, "4. Critical Analysis")
    add_heading(document, "Strengths", level=2)
    add_paragraph(
        document,
        "The selected GRU model has several strengths. Its main advantage is that it can learn sequential patterns "
        "from historical stock prices while using a simpler gated structure than LSTM. The update gate allows the "
        "model to keep useful past information, while the reset gate helps it decide when older information should "
        "be ignored. This is useful for stock data because recent price movement is important, but older trends may "
        "still contain useful context.",
    )
    add_paragraph(
        document,
        "The numerical results support the choice of GRU as the strongest RNN-based model. GRU achieved RMSE = "
        "26.06, MAE = 15.54, MAPE = 6.28%, and R2 = 0.695. Compared with the vanilla RNN and LSTM, GRU produced "
        "lower RMSE, MAE, and MAPE, showing that its update and reset gates handled the 60-day sequence more "
        "effectively. GRU also achieved lower MAE and MAPE than the 60-day moving average baseline, although the "
        "baseline still achieved a lower RMSE and higher R2.",
    )
    add_paragraph(
        document,
        "Another strength of the project design is that it reduces data leakage. The train/validation/test split is "
        "chronological rather than random, which is more suitable for time-series forecasting because the model "
        "should only learn from past data and then predict future data. The MinMaxScaler was fitted only on the "
        "training set before being applied to the validation and test sets, so information from later periods was "
        "not used during normalization. The validation set was used to monitor generalization and select the best "
        "epoch before final testing.",
    )
    add_heading(document, "Limitations", level=2)
    add_paragraph(
        document,
        "However, the model has important limitations. First, there is still a risk of overfitting. RNN-based models "
        "contain many trainable parameters, and stock price data is noisy. A model may learn the shape of the training "
        "period well but fail when the market enters a different condition. Validation loss reduces this risk by "
        "showing whether performance improves on unseen validation data, but it does not guarantee reliable "
        "performance in future market regimes.",
    )
    add_paragraph(
        document,
        "Second, the model is highly dependent on historical price data. This project uses only the adjusted closing "
        "price as the input feature. In real financial markets, stock prices are affected by earnings announcements, "
        "interest rates, inflation, sector performance, market indexes, regulation, product news, and investor "
        "sentiment. Since these variables are not included, the model cannot directly respond to important information "
        "outside the price series.",
    )
    add_paragraph(
        document,
        "Third, the model is sensitive to market volatility and sudden events. Unexpected events such as financial "
        "crises, product announcements, lawsuits, geopolitical shocks, or earnings surprises can change prices quickly. "
        "Because the GRU learns from historical patterns, it may react slowly to sudden turning points. Strong test "
        "performance on historical data therefore does not guarantee reliable prediction during unusual market conditions.",
    )
    add_paragraph(
        document,
        "There are also evaluation limitations. The model was tested using one stock, one date range, and one "
        "chronological split. A stronger evaluation would use rolling-window validation, where the model is trained "
        "and tested across multiple time periods. This would show whether the model remains stable across different "
        "market conditions. Longer forecasting horizons would also be more difficult and should be tested separately.",
    )
    add_paragraph(
        document,
        "Future improvements could include adding trading volume, technical indicators, market index prices, interest "
        "rates, and sentiment features. Regularization methods such as early stopping, validation loss monitoring, "
        "and hyperparameter tuning could reduce overfitting risk. A Transformer-based time-series model could also "
        "be tested. Overall, the GRU model is effective for this experiment, but it should be viewed as a forecasting "
        "aid rather than a guaranteed predictor of future stock prices.",
    )
    add_heading(document, "Comparison with Alternative Models", level=2)
    add_paragraph(
        document,
        "The comparison with alternative models shows a trade-off between accuracy, simplicity, and evaluation "
        "metric. The moving average baseline is very easy to understand and does not require training; in this run "
        "it achieved the best RMSE and R2, showing that a simple baseline can remain strong for one-step stock "
        "prediction. GRU was the best neural model and achieved the best MAE and MAPE, meaning its average absolute "
        "prediction error was lower. The vanilla RNN was simpler than GRU and LSTM, but it performed worse because "
        "it has weaker memory handling. LSTM is more powerful in theory because it has a separate memory cell and "
        "three gates, but in this experiment its extra complexity did not improve performance. Therefore, GRU is "
        "the best RNN-based choice for this dataset, while the moving average baseline remains an important benchmark.",
    )

    add_heading(document, "Conclusion")
    add_paragraph(
        document,
        "This project designed and applied multiple RNN-based models for stock price prediction using public GOOGL "
        "historical financial data. The data was cleaned, normalized, split chronologically into training, validation, "
        "and testing sets, converted into 60-day sequences, and evaluated using RMSE, MAE, MAPE, and R2. GRU was "
        "the strongest RNN-based model and achieved the lowest MAE and MAPE, while the moving average baseline "
        "achieved the best RMSE and R2. This shows that recurrent neural networks can learn useful time-dependent "
        "patterns from historical stock data, but they must still be compared against simple baselines. The model "
        "remains limited by overfitting risk, historical-data dependence, and unpredictable market volatility.",
    )

    add_heading(document, "Source Code")
    add_source_code_link(document)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
